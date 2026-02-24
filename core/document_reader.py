"""
Document Reader

Extracts text content from any document format.
Supports: PDF, DOCX, XLSX, XML, TXT, and more.
"""

import os
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional
from datetime import datetime

from .logging_config import get_logger

logger = get_logger("document_reader")


@dataclass
class DocumentContent:
    """Extracted content from a document."""
    path: str
    filename: str
    extension: str
    text: str
    raw_content: Optional[bytes] = None
    metadata: dict = field(default_factory=dict)
    extraction_method: str = "unknown"
    extraction_time: datetime = field(default_factory=datetime.now)
    error: Optional[str] = None

    @property
    def is_successful(self) -> bool:
        return self.error is None and len(self.text.strip()) > 0


class DocumentReader:
    """
    Reads documents of any format and extracts text content.

    This is a best-effort extractor - it tries multiple methods
    and returns whatever text can be extracted.
    """

    SUPPORTED_EXTENSIONS = {
        '.pdf', '.docx', '.xlsx', '.xml', '.txt', '.csv',
        '.html', '.htm', '.md', '.json', '.x83', '.d83',
        '.xsl', '.aidocdef', '.aidoc', '.aiform',
    }

    def __init__(self):
        logger.debug("Initializing DocumentReader")
        self._pdf_available = self._check_pdf_support()
        self._docx_available = self._check_docx_support()
        self._xlsx_available = self._check_xlsx_support()
        logger.debug(f"Library support: PDF={self._pdf_available}, DOCX={self._docx_available}, XLSX={self._xlsx_available}")

    def _check_pdf_support(self) -> bool:
        try:
            import pypdf
            logger.debug("pypdf library available")
            return True
        except ImportError:
            try:
                import PyPDF2
                logger.debug("PyPDF2 library available")
                return True
            except ImportError:
                logger.warning("No PDF library available (pypdf or PyPDF2)")
                return False

    def _check_docx_support(self) -> bool:
        try:
            import docx
            logger.debug("python-docx library available")
            return True
        except ImportError:
            logger.warning("python-docx library not available")
            return False

    def _check_xlsx_support(self) -> bool:
        try:
            import openpyxl
            logger.debug("openpyxl library available")
            return True
        except ImportError:
            logger.warning("openpyxl library not available")
            return False

    def read(self, path: str) -> DocumentContent:
        """Read a document and extract its text content."""
        path = os.path.abspath(path)
        filename = os.path.basename(path)
        extension = os.path.splitext(filename)[1].lower()

        logger.debug(f"Reading document: {filename} (extension: {extension})")

        if not os.path.exists(path):
            logger.error(f"File not found: {path}")
            return DocumentContent(
                path=path,
                filename=filename,
                extension=extension,
                text="",
                error=f"File not found: {path}"
            )

        try:
            if extension == '.pdf':
                result = self._read_pdf(path)
            elif extension == '.docx':
                result = self._read_docx(path)
            elif extension == '.xlsx':
                result = self._read_xlsx(path)
            elif extension in ('.xml', '.xsl', '.aidocdef', '.aidoc', '.aiform'):
                result = self._read_xml(path)
            elif extension in ('.txt', '.md', '.csv', '.json'):
                result = self._read_text(path)
            elif extension in ('.html', '.htm'):
                result = self._read_html(path)
            elif extension in ('.x83', '.d83'):
                result = self._read_gaeb(path)
            else:
                logger.debug(f"Unknown extension {extension}, attempting generic read")
                result = self._read_unknown(path)

            if result.is_successful:
                logger.info(f"Successfully read {filename}: {len(result.text)} chars via {result.extraction_method}")
            else:
                logger.warning(f"Failed to read {filename}: {result.error}")

            return result

        except Exception as e:
            logger.exception(f"Unexpected error reading {filename}")
            return DocumentContent(
                path=path,
                filename=filename,
                extension=extension,
                text="",
                error=f"Error reading file: {str(e)}"
            )

    def _read_pdf(self, path: str) -> DocumentContent:
        """Extract text from PDF."""
        filename = os.path.basename(path)
        extension = '.pdf'
        logger.debug(f"Attempting PDF extraction for {filename}")

        try:
            import pypdf
            logger.debug("Using pypdf for extraction")
            reader = pypdf.PdfReader(path)
            text_parts = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(f"--- Page {i+1} ---\n{page_text}")
                logger.debug(f"Page {i+1}: extracted {len(page_text)} chars")

            return DocumentContent(
                path=path,
                filename=filename,
                extension=extension,
                text="\n\n".join(text_parts),
                extraction_method="pypdf",
                metadata={"page_count": len(reader.pages)}
            )
        except ImportError:
            logger.debug("pypdf not available, trying PyPDF2")

        try:
            import PyPDF2
            logger.debug("Using PyPDF2 for extraction")
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text_parts = []
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_parts.append(f"--- Page {i+1} ---\n{page_text}")

                return DocumentContent(
                    path=path,
                    filename=filename,
                    extension=extension,
                    text="\n\n".join(text_parts),
                    extraction_method="PyPDF2",
                    metadata={"page_count": len(reader.pages)}
                )
        except ImportError:
            logger.error("No PDF library available")

        return DocumentContent(
            path=path,
            filename=filename,
            extension=extension,
            text="",
            error="No PDF library available. Install pypdf: pip install pypdf"
        )

    def _read_docx(self, path: str) -> DocumentContent:
        """Extract text from DOCX."""
        filename = os.path.basename(path)
        extension = '.docx'
        logger.debug(f"Attempting DOCX extraction for {filename}")

        try:
            import docx
            logger.debug("Using python-docx for extraction")
            doc = docx.Document(path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract text from tables
            table_count = len(doc.tables)
            logger.debug(f"Found {len(doc.paragraphs)} paragraphs and {table_count} tables")

            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_rows.append(row_text)
                if table_rows:
                    text_parts.append("\n[TABLE]\n" + "\n".join(table_rows) + "\n[/TABLE]")

            return DocumentContent(
                path=path,
                filename=filename,
                extension=extension,
                text="\n\n".join(text_parts),
                extraction_method="python-docx",
            )
        except ImportError:
            logger.debug("python-docx not available, trying ZIP fallback")
            return self._read_docx_as_zip(path)

    def _read_docx_as_zip(self, path: str) -> DocumentContent:
        """Extract text from DOCX by treating it as a ZIP file."""
        filename = os.path.basename(path)
        extension = '.docx'
        logger.debug(f"Using ZIP fallback for DOCX: {filename}")

        try:
            with zipfile.ZipFile(path, 'r') as z:
                with z.open('word/document.xml') as f:
                    tree = ET.parse(f)
                    root = tree.getroot()

                    text_parts = []
                    for elem in root.iter():
                        if elem.tag.endswith('}t'):
                            if elem.text:
                                text_parts.append(elem.text)

                    return DocumentContent(
                        path=path,
                        filename=filename,
                        extension=extension,
                        text=" ".join(text_parts),
                        extraction_method="docx-as-zip",
                    )
        except Exception as e:
            logger.error(f"ZIP fallback failed for {filename}: {e}")
            return DocumentContent(
                path=path,
                filename=filename,
                extension=extension,
                text="",
                error=f"Could not read DOCX: {str(e)}. Install python-docx: pip install python-docx"
            )

    def _read_xlsx(self, path: str) -> DocumentContent:
        """Extract text from XLSX."""
        filename = os.path.basename(path)
        extension = '.xlsx'
        logger.debug(f"Attempting XLSX extraction for {filename}")

        try:
            import openpyxl
            logger.debug("Using openpyxl for extraction")
            wb = openpyxl.load_workbook(path, data_only=True)
            text_parts = []

            logger.debug(f"Workbook has {len(wb.sheetnames)} sheets: {wb.sheetnames}")

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_text = [f"=== Sheet: {sheet_name} ==="]

                for row in sheet.iter_rows():
                    row_values = []
                    for cell in row:
                        if cell.value is not None:
                            row_values.append(str(cell.value))
                    if row_values:
                        sheet_text.append(" | ".join(row_values))

                if len(sheet_text) > 1:
                    text_parts.append("\n".join(sheet_text))

            return DocumentContent(
                path=path,
                filename=filename,
                extension=extension,
                text="\n\n".join(text_parts),
                extraction_method="openpyxl",
                metadata={"sheet_count": len(wb.sheetnames)}
            )
        except ImportError:
            logger.error("openpyxl not available")
            return DocumentContent(
                path=path,
                filename=filename,
                extension=extension,
                text="",
                error="No XLSX library available. Install openpyxl: pip install openpyxl"
            )

    def _read_xml(self, path: str) -> DocumentContent:
        """Extract text from XML."""
        filename = os.path.basename(path)
        extension = '.xml'
        logger.debug(f"Attempting XML extraction for {filename}")

        try:
            # Try different encodings
            content = None
            used_encoding = None
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    with open(path, 'r', encoding=encoding) as f:
                        content = f.read()
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue

            if content is None:
                with open(path, 'rb') as f:
                    content = f.read().decode('utf-8', errors='replace')
                used_encoding = 'utf-8-fallback'

            logger.debug(f"Read XML with encoding: {used_encoding}")

            tree = ET.fromstring(content)

            def extract_text(elem, depth=0):
                text_parts = []
                tag_name = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag

                if elem.text and elem.text.strip():
                    text_parts.append(f"{tag_name}: {elem.text.strip()}")

                for attr, value in elem.attrib.items():
                    text_parts.append(f"{tag_name}@{attr}: {value}")

                for child in elem:
                    text_parts.extend(extract_text(child, depth + 1))

                return text_parts

            text = "\n".join(extract_text(tree))

            return DocumentContent(
                path=path,
                filename=filename,
                extension=extension,
                text=text,
                extraction_method="xml-parser",
            )

        except ET.ParseError as e:
            logger.warning(f"XML parsing failed for {filename}: {e}, falling back to raw text")
            with open(path, 'r', encoding='utf-8', errors='replace') as f:
                return DocumentContent(
                    path=path,
                    filename=filename,
                    extension=extension,
                    text=f.read(),
                    extraction_method="raw-text",
                )

    def _read_text(self, path: str) -> DocumentContent:
        """Read plain text file."""
        filename = os.path.basename(path)
        extension = os.path.splitext(filename)[1].lower()
        logger.debug(f"Attempting text extraction for {filename}")

        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    content = f.read()
                logger.debug(f"Read text file with encoding: {encoding}")
                return DocumentContent(
                    path=path,
                    filename=filename,
                    extension=extension,
                    text=content,
                    extraction_method=f"text-{encoding}",
                )
            except UnicodeDecodeError:
                continue

        logger.debug("Falling back to utf-8 with error replacement")
        with open(path, 'r', encoding='utf-8', errors='replace') as f:
            return DocumentContent(
                path=path,
                filename=filename,
                extension=extension,
                text=f.read(),
                extraction_method="text-fallback",
            )

    def _read_html(self, path: str) -> DocumentContent:
        """Extract text from HTML."""
        filename = os.path.basename(path)
        extension = os.path.splitext(filename)[1].lower()
        logger.debug(f"Attempting HTML extraction for {filename}")

        with open(path, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()

        import re
        content = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL | re.IGNORECASE)
        content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL | re.IGNORECASE)
        content = re.sub(r'<[^>]+>', ' ', content)
        content = re.sub(r'\s+', ' ', content)

        return DocumentContent(
            path=path,
            filename=filename,
            extension=extension,
            text=content.strip(),
            extraction_method="html-strip",
        )

    def _read_gaeb(self, path: str) -> DocumentContent:
        """Read GAEB files (German construction tendering format)."""
        logger.debug(f"Reading GAEB file as XML: {os.path.basename(path)}")
        return self._read_xml(path)

    def _read_unknown(self, path: str) -> DocumentContent:
        """Attempt to read an unknown file format."""
        filename = os.path.basename(path)
        extension = os.path.splitext(filename)[1].lower()
        logger.debug(f"Attempting generic read for unknown format: {extension}")

        try:
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            return DocumentContent(
                path=path,
                filename=filename,
                extension=extension,
                text=content,
                extraction_method="text-attempt",
            )
        except UnicodeDecodeError:
            pass

        for encoding in ['latin-1', 'cp1252']:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    content = f.read()
                logger.debug(f"Read unknown format with encoding: {encoding}")
                return DocumentContent(
                    path=path,
                    filename=filename,
                    extension=extension,
                    text=content,
                    extraction_method=f"text-{encoding}",
                )
            except:
                continue

        logger.warning(f"Unable to extract text from {extension} file: {filename}")
        return DocumentContent(
            path=path,
            filename=filename,
            extension=extension,
            text="",
            error=f"Unable to extract text from {extension} file"
        )

    def read_directory(self, directory: str, recursive: bool = True) -> list[DocumentContent]:
        """Read all documents in a directory."""
        documents = []
        directory = Path(directory)

        logger.info(f"Scanning directory: {directory} (recursive={recursive})")

        if recursive:
            files = list(directory.rglob('*'))
        else:
            files = list(directory.glob('*'))

        logger.debug(f"Found {len(files)} total files")

        skipped = {"hidden": 0, "archive": 0, "image": 0, "state": 0}

        for file_path in files:
            if file_path.is_file():
                if file_path.name.startswith('.'):
                    skipped["hidden"] += 1
                    continue
                if file_path.suffix.lower() in ('.zip', '.rar', '.7z', '.tar', '.gz'):
                    skipped["archive"] += 1
                    continue
                if file_path.suffix.lower() in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.ico'):
                    skipped["image"] += 1
                    continue
                # Skip state directory files (logs, saved graphs, etc.)
                if '.tender_state' in str(file_path):
                    skipped["state"] += 1
                    continue

                doc = self.read(str(file_path))
                documents.append(doc)

        logger.info(f"Read {len(documents)} documents, skipped: {skipped}")
        return documents

    def extract_archives(self, directory: str) -> list[str]:
        """Extract all ZIP files in a directory. Returns list of extracted directories."""
        extracted = []
        directory = Path(directory)

        zip_files = list(directory.rglob('*.zip'))
        logger.info(f"Found {len(zip_files)} ZIP archives to extract")

        for zip_path in zip_files:
            extract_dir = zip_path.with_suffix('')
            if not extract_dir.exists():
                try:
                    import subprocess
                    import platform

                    logger.debug(f"Extracting: {zip_path.name} -> {extract_dir}")

                    if platform.system() == 'Darwin':
                        subprocess.run(
                            ['ditto', '-x', '-k', str(zip_path), str(extract_dir)],
                            check=True,
                            capture_output=True
                        )
                    else:
                        with zipfile.ZipFile(zip_path, 'r') as z:
                            z.extractall(extract_dir)

                    extracted.append(str(extract_dir))
                    logger.info(f"Extracted: {zip_path.name}")
                except Exception as e:
                    logger.error(f"Failed to extract {zip_path.name}: {e}")
            else:
                logger.debug(f"Already extracted: {zip_path.name}")

        return extracted


def read_document(path: str) -> DocumentContent:
    """Read a single document."""
    reader = DocumentReader()
    return reader.read(path)


def read_all_documents(directory: str, extract_archives: bool = True) -> list[DocumentContent]:
    """Read all documents in a directory, optionally extracting archives first."""
    reader = DocumentReader()

    if extract_archives:
        reader.extract_archives(directory)

    return reader.read_directory(directory)
